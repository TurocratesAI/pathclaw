#!/usr/bin/env python3
"""Build docs/PathClaw-Manual.docx from MANUAL.md.

Prefers pandoc if available. Falls back to python-docx paragraph rendering.

Usage:
    python scripts/build_manual.py
"""
from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
MANUAL_MD = ROOT / "MANUAL.md"
DOCX_OUT = ROOT / "docs/PathClaw-Manual.docx"


def build_with_pandoc() -> bool:
    if not shutil.which("pandoc"):
        return False
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    subprocess.check_call([
        "pandoc", str(MANUAL_MD),
        "-f", "gfm",
        "-o", str(DOCX_OUT),
        "--toc",
        "--metadata", "title=PathClaw Manual",
    ])
    return True


def build_with_python_docx() -> bool:
    try:
        from docx import Document
    except ImportError:
        print("pandoc not found and python-docx not installed.", file=sys.stderr)
        print("    pip install python-docx     OR     apt install pandoc", file=sys.stderr)
        return False

    doc = Document()
    doc.add_heading("PathClaw Manual", level=0)
    for raw in MANUAL_MD.read_text().splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("```"):
            doc.add_paragraph(line, style="Intense Quote")
        else:
            doc.add_paragraph(line)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    return True


def main() -> int:
    if not MANUAL_MD.exists():
        print(f"MANUAL.md not found at {MANUAL_MD}", file=sys.stderr)
        return 1
    if build_with_pandoc():
        print(f"Built {DOCX_OUT} via pandoc")
        return 0
    if build_with_python_docx():
        print(f"Built {DOCX_OUT} via python-docx fallback")
        return 0
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
